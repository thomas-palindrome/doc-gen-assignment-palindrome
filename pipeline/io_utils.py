"""Loading client source material and preparing it for the generator."""

import json
from pathlib import Path

from docx import Document


def read_docx(path: Path) -> str:
    """Return the text of a .docx as plain text, one paragraph per line."""
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def load_client(client_dir: Path) -> dict:
    """Load every source for a client into a dict."""
    db_path = client_dir / "client_data_db.json"
    return {
        "meeting_notes": read_docx(client_dir / "meeting_notes.docx"),
        "report_request": read_docx(client_dir / "report_request.docx"),
        "client_data_db": json.loads(read_text(db_path)),
        "fde_notes": read_text(client_dir / "fde_notes.md"),
        "template_spec": read_text(client_dir / "template_spec.md"),
    }


def build_context(client: dict) -> str:
    """Assemble all of the client's source material into a single context string
    that gets passed to the model alongside each section's instruction."""
    return "\n\n".join(
        [
            "=== MEETING NOTES ===",
            client["meeting_notes"],
            "=== REPORT REQUEST ===",
            client["report_request"],
            "=== CLIENT DATA (DB) ===",
            json.dumps(client["client_data_db"], indent=2),
            "=== INTERNAL NOTES ===",
            client["fde_notes"],
            "=== TEMPLATE SPEC ===",
            client["template_spec"],
        ],
    )
