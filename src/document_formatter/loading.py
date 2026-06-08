"""Reading a client's source files into plain text for the generator."""

from pathlib import Path

from docx import Document


def read_docx(path: Path) -> str:
    """Return the text of a .docx as plain text, one paragraph per line, tables pipe-joined."""
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def read_file(path: Path) -> str:
    """Read one source file as text, dispatched by file type."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx(path)
    if suffix in {".png", ".jpg", ".jpeg"}:
        return f"[image file, not read as text: {path.name}]"
    return read_text(path)
